# -*- coding: utf-8 -*-
"""마크다운 -> DOCX 변환 (제목/표/목록/코드블록/인용/굵게 지원, 한글 폰트 적용)"""
import io, re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

KFONT = "맑은 고딕"
MONO = "D2Coding"

def set_font(run, name=KFONT, size=10, bold=False, color=None, mono=False):
    f = name if not mono else MONO
    run.font.name = f
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.append(rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rf.set(qn(a), f)

INLINE = re.compile(r'(\*\*.+?\*\*|`[^`]+`)')

def add_inline(p, text, size=10, base_bold=False):
    """**굵게** 와 `코드` 를 해석해 run 으로 추가"""
    text = text.replace('<br>', '\n')
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            set_font(p.add_run(part[2:-2]), size=size, bold=True)
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            r = p.add_run(part[1:-1])
            set_font(r, size=size - 0.5, mono=True, color=(0xB0, 0x30, 0x60))
        else:
            set_font(p.add_run(part), size=size, bold=base_bold)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def convert(md_path, out_path):
    lines = io.open(md_path, encoding='utf-8').read().split('\n')
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = KFONT; st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), KFONT)
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.0)
        s.top_margin = s.bottom_margin = Cm(2.0)

    i = 0
    while i < len(lines):
        ln = lines[i]; s = ln.strip()

        # 코드블록
        if s.startswith('```'):
            i += 1; buf = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                buf.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph(); pf = p.paragraph_format
            pf.left_indent = Cm(0.5); pf.space_before = Pt(4); pf.space_after = Pt(6)
            pf.line_spacing = 1.0
            set_font(p.add_run('\n'.join(buf)), size=8.5, mono=True, color=(0x22, 0x22, 0x22))
            shade_p = p._p.get_or_add_pPr()
            sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), 'F2F3F5')
            shade_p.append(sh)
            continue

        # 표
        if s.startswith('|') and i + 1 < len(lines) and re.match(r'^\s*\|[\s:|-]+\|\s*$', lines[i + 1]):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                if not re.match(r'^\s*\|[\s:|-]+\|\s*$', lines[i]):
                    rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')])
                i += 1
            ncol = max(len(r) for r in rows)
            t = doc.add_table(rows=0, cols=ncol)
            t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                cells = t.add_row().cells
                for ci in range(ncol):
                    txt = row[ci] if ci < len(row) else ''
                    cp = cells[ci].paragraphs[0]
                    cp.paragraph_format.space_before = Pt(2)
                    cp.paragraph_format.space_after = Pt(2)
                    if ri == 0:
                        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        shade(cells[ci], 'DCE6F1')
                    add_inline(cp, txt, size=9, base_bold=(ri == 0))
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        # 구분선
        if s == '---':
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
            bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
            bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), 'AAAAAA')
            pbdr.append(bot); pPr.append(pbdr)
            i += 1; continue

        # 제목
        m = re.match(r'^(#{1,4})\s+(.*)', s)
        if m:
            lvl = len(m.group(1)); txt = m.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt({1: 0, 2: 14, 3: 10, 4: 8}[lvl])
            p.paragraph_format.space_after = Pt(6)
            size = {1: 16, 2: 13, 3: 11.5, 4: 10.5}[lvl]
            col = {1: (0x1F, 0x38, 0x64), 2: (0x1F, 0x38, 0x64), 3: (0x2E, 0x4A, 0x7A), 4: (0x33, 0x33, 0x33)}[lvl]
            if lvl == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for part in INLINE.split(txt):
                if not part: continue
                if part.startswith('**') and part.endswith('**'):
                    part = part[2:-2]
                set_font(p.add_run(part), size=size, bold=True, color=col)
            i += 1; continue

        # 인용
        if s.startswith('>'):
            buf = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                buf.append(lines[i].strip().lstrip('>').strip()); i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
            add_inline(p, '\n'.join(x for x in buf if x), size=9.5)
            for r in p.runs:
                r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            continue

        # 목록
        m = re.match(r'^(\s*)[-*]\s+(.*)', ln)
        if m:
            depth = len(m.group(1)) // 2
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5 + 0.5 * depth)
            p.paragraph_format.space_after = Pt(2)
            set_font(p.add_run('• ' if depth == 0 else '- '), size=10)
            add_inline(p, m.group(2))
            i += 1; continue

        m = re.match(r'^(\s*)(\d+)\.\s+(.*)', ln)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(2)
            set_font(p.add_run(m.group(2) + '. '), size=10, bold=True)
            add_inline(p, m.group(3))
            i += 1; continue

        if not s:
            i += 1; continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_inline(p, s)
        i += 1

    doc.save(out_path)
    print("OK ->", out_path)

if __name__ == '__main__':
    convert(sys.argv[1], sys.argv[2])
